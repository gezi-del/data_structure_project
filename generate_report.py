import contextlib
import io
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from experiment_runner import run_all_experiments


REPORT_PATH = Path("新能源物流车队调度项目报告.docx")


def set_run_font(run, size=None, bold=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text))
    set_run_font(run, size=10, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_text(hdr[idx], header, bold=True)
        shade_cell(hdr[idx], "F2F4F7")
        if widths:
            hdr[idx].width = Inches(widths[idx])
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
            if widths:
                cells[idx].width = Inches(widths[idx])
    doc.add_paragraph()
    return table


def add_heading(doc, text, level):
    paragraph = doc.add_heading(level=level)
    run = paragraph.add_run(text)
    set_run_font(run, size=16 if level == 1 else 13, bold=True, color="2E74B5")
    return paragraph


def add_body(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.1
    run = paragraph.add_run(text)
    set_run_font(run, size=11)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    set_run_font(run, size=11)
    return paragraph


def build_report():
    with contextlib.redirect_stdout(io.StringIO()):
        experiments = run_all_experiments(force=True)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("新能源物流车队协同调度系统项目报告")
    set_run_font(run, size=22, bold=True, color="0B2545")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("数据结构大作业 | 图结构寻路、动态调度、充电排队与遗传算法离线优化")
    set_run_font(run, size=11, color="555555")
    doc.add_paragraph()

    add_heading(doc, "1. 题目要求对照", 1)
    add_table(
        doc,
        ["要求", "实现情况", "对应模块"],
        [
            ["图结构实现道路和寻路", "随机生成全连通带权无向图，并用 Dijkstra 求最短路径。", "Graph.py, Fleet_Controller.py"],
            ["有限车辆、电量、载重", "每辆车维护电量上限、当前电量、载重上限和当前载重。", "Vehicle.py"],
            ["动态任务", "任务按概率或预设时间表动态出现，包含时间、坐标和货物重量。", "Task_Manager.py, Task.py"],
            ["收益评分与超时扣分", "完成时间越短、路径越短得分越高，超时额外扣分。", "Fleet_Controller.py"],
            ["电量不足补能", "车辆判断安全余量，不足时选择综合旅行时间和排队等待最小的充电站。", "Fleet_Controller.py, Graph.py"],
            ["充电站排队负载", "每个充电站维护 queue_count、queue_limit 和满负荷状态。", "Graph.py"],
            ["多车协同", "大重量任务可拆分，由多辆车分别运输部分载重。", "Fleet_Controller.py"],
            ["至少两种调度策略", "实现最近任务优先和最大重量优先。", "Fleet_Controller.py"],
            ["三种规模模拟", "支持小、中、大三种问题规模并自动实验对比。", "simulation_config.py, experiment_runner.py"],
            ["进阶算法", "实现遗传算法离线优化，用于上帝视角近似全局方案对比。", "genetic_optimizer.py"],
            ["图形界面展示", "实现调度驾驶舱，展示地图、车辆、任务、充电负载和实验对比。", "templates/index.html"],
        ],
        widths=[1.45, 3.35, 1.55],
    )

    add_heading(doc, "2. 系统总体设计", 1)
    add_body(
        doc,
        "系统采用 Flask 后端和 Canvas 前端。后端负责地图生成、任务生成、车辆状态推进、在线调度策略、遗传算法实验和 API 输出；前端负责实时可视化和交互控制。"
    )
    for item in [
        "地图层：生成全连通带权无向图，节点可标记为仓库、充电站、普通道路或任务点。",
        "调度层：每个时间步先生成任务，再更新充电、车辆移动、任务分配和返仓逻辑。",
        "展示层：页面可切换规模、策略和多车协同开关，并显示收益、完成数、超时数、里程和充电等待。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "3. 数据结构与算法", 1)
    add_body(
        doc,
        "道路网络使用邻接表保存：nodes 保存节点对象，adj 保存边权。邻接表适合稀疏道路网络，查询相邻道路和执行 Dijkstra 都较高效。"
    )
    add_body(
        doc,
        "最短路使用 Dijkstra，并通过 ShortestPathCache 缓存起终点路径，避免调度过程中重复计算。充电站选择不是简单最近站，而是综合到站时间和当前排队等待时间。"
    )
    add_table(
        doc,
        ["对象", "关键字段", "作用"],
        [
            ["Vehicle", "battery_capacity, current_battery, load_capacity, current_load", "描述新能源车的电量与载重约束。"],
            ["Task", "appear_time, target_id, coords, weight, status, score", "描述配送任务和评分结果。"],
            ["Node", "type, queue_count, queue_limit, is_full", "描述道路节点及充电站负载。"],
            ["VehicleState", "route, mission, task_id, wait_time_left, charge_time_left", "描述车辆在调度器中的任务状态。"],
        ],
        widths=[1.2, 2.9, 2.25],
    )

    add_heading(doc, "4. 调度策略", 1)
    add_body(doc, "在线调度策略只知道当前已经出现的任务，符合真实动态调度场景。")
    add_bullet(doc, "最近任务优先：空闲车辆从候选任务中选择最短路径距离最近的任务，目标是减少单次响应距离。")
    add_bullet(doc, "最大重量优先：优先选择货物重量大的任务，目标是尽快处理高收益或高压力订单。")
    add_bullet(doc, "多车协同：当任务重量超过单车载重时，将剩余重量分批分配给多辆车，直到任务完成。")

    add_heading(doc, "5. 遗传算法离线优化", 1)
    add_body(
        doc,
        "遗传算法用于离线对比：假设一段时间内任务全部已知，把任务配送顺序和车辆分配编码为染色体。适应度函数综合总收益、完成任务数、超时罚分、总里程和充电等待时间。"
    )
    add_table(
        doc,
        ["环节", "设计"],
        [
            ["染色体", "任务配送单元顺序 + 每个配送单元分配的车辆编号。"],
            ["初始种群", "随机解与按出现时间、距离、重量构造的启发式解混合。"],
            ["选择", "锦标赛选择，优先保留适应度高的方案。"],
            ["交叉", "顺序交叉保持任务序列合法，车辆分配采用均匀交叉。"],
            ["变异", "随机交换配送顺序，或随机改变部分车辆分配。"],
            ["精英保留", "每代保留当前最优个体，保证最优适应度不下降。"],
        ],
        widths=[1.4, 4.95],
    )

    add_heading(doc, "6. 实验结果", 1)
    rows = []
    for group in experiments["results"]:
        for strategy in group["strategies"]:
            rows.append([
                group["scale_label"],
                strategy["label"],
                strategy["total_score"],
                strategy["completed_tasks"],
                strategy["timeout_tasks"],
                strategy["avg_finish_time"],
                strategy["total_distance"],
            ])
    add_table(
        doc,
        ["规模", "策略", "总收益", "完成数", "超时数", "平均完成(s)", "累计里程"],
        rows,
        widths=[0.8, 1.25, 0.9, 0.75, 0.75, 1.0, 0.9],
    )
    add_body(
        doc,
        "实验显示，遗传算法在小规模和中规模中获得更高收益，并在大规模中完成更多任务；在线策略实时性更强，但只能基于已出现任务做局部决策。"
    )

    add_heading(doc, "7. 图形界面说明", 1)
    add_body(
        doc,
        "前端改为调度驾驶舱布局：左侧控制规模、策略和运行状态，中间显示地图与车辆动画，右侧显示任务、车辆、充电站负载和策略实验表。该界面适合课堂演示视频录制，也方便评分人员快速看到功能覆盖情况。"
    )

    add_heading(doc, "8. 总结与不足", 1)
    add_body(
        doc,
        "本项目完整覆盖了新能源物流车队调度题目的核心要求，并额外加入遗传算法离线优化作为进阶算法。当前遗传算法仍是近似求解，未接入精确求解器；后续可继续加入更细致的电池衰减、道路拥堵、时间窗和真实地图数据。"
    )

    doc.save(REPORT_PATH)
    return REPORT_PATH


if __name__ == "__main__":
    print(build_report())
